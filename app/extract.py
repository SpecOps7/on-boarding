"""Per-filetype text extraction from synced Box files."""

from dataclasses import dataclass
from pathlib import Path

MIN_PAGE_CHARS = 40  # a PDF page yielding less than this is treated as scanned

TEXT_PLAIN_EXTS = {".txt", ".md", ".csv", ".log"}
IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tif", ".tiff", ".webp"}
LEGACY_EXTS = {".doc", ".xls", ".ppt"}


@dataclass
class Section:
    loc: str
    text: str


@dataclass
class Extraction:
    # status: indexed | attachable | unsupported_legacy | not_indexed | error
    status: str
    sections: list[Section]
    note: str = ""


def extract(path: Path) -> Extraction:
    ext = path.suffix.lower()
    try:
        if ext == ".pdf":
            return _extract_pdf(path)
        if ext == ".docx":
            return _extract_docx(path)
        if ext == ".xlsx":
            return _extract_xlsx(path)
        if ext in TEXT_PLAIN_EXTS:
            text = path.read_text(errors="replace")[:500_000]
            return Extraction("indexed", [Section("file", text)]) if text.strip() \
                else Extraction("not_indexed", [], "empty file")
        if ext in IMAGE_EXTS:
            # No OCR: index the filename so questions can match it; Claude can
            # open the image itself at answer time via the Read tool.
            return Extraction("attachable", [Section("file", _name_tokens(path))], "image")
        if ext in LEGACY_EXTS:
            return Extraction("unsupported_legacy", [], "legacy Office format — open in Box")
        return Extraction("not_indexed", [], f"unsupported type {ext or '(none)'}")
    except Exception as e:  # noqa: BLE001 - one bad file must not sink the folder
        return Extraction("error", [], f"{type(e).__name__}: {e}"[:300])


def _name_tokens(path: Path) -> str:
    import re
    return " ".join(re.findall(r"[A-Za-z0-9]+", str(path.name))) + " " + \
        " ".join(re.findall(r"[A-Za-z0-9]+", str(path.parent.name)))


def _extract_pdf(path: Path) -> Extraction:
    from pypdf import PdfReader

    reader = PdfReader(path)
    sections, scanned_pages = [], 0
    for i, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if len(text) >= MIN_PAGE_CHARS:
            sections.append(Section(f"p.{i}", text))
        else:
            scanned_pages += 1
    if not sections:
        # Fully scanned PDF: searchable by filename, readable by Claude directly.
        return Extraction("attachable", [Section("file", _name_tokens(path))],
                          f"scanned PDF ({scanned_pages} image pages)")
    note = f"{scanned_pages} scanned pages skipped" if scanned_pages else ""
    return Extraction("indexed", sections, note)


def _extract_docx(path: Path) -> Extraction:
    import docx

    doc = docx.Document(path)
    blocks: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            blocks.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                blocks.append("\t".join(cells))
    if not blocks:
        return Extraction("not_indexed", [], "no extractable text")

    # Group blocks into ~3000-char sections; the chunker splits further.
    sections, buf, n = [], [], 1
    size = 0
    for b in blocks:
        buf.append(b)
        size += len(b)
        if size >= 3000:
            sections.append(Section(f"§{n}", "\n".join(buf)))
            buf, size, n = [], 0, n + 1
    if buf:
        sections.append(Section(f"§{n}", "\n".join(buf)))
    return Extraction("indexed", sections)


def _extract_xlsx(path: Path) -> Extraction:
    import openpyxl

    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    sections = []
    try:
        for ws in wb.worksheets:
            rows = []
            for r, row in enumerate(ws.iter_rows(values_only=True)):
                if r >= 500:
                    rows.append("… (truncated at 500 rows)")
                    break
                cells = [str(c) for c in row if c is not None and str(c).strip()]
                if cells:
                    rows.append("\t".join(cells))
            if rows:
                sections.append(Section(f"sheet:{ws.title}", "\n".join(rows)))
    finally:
        wb.close()
    if not sections:
        return Extraction("not_indexed", [], "no extractable cells")
    return Extraction("indexed", sections)
